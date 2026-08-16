import io
import psycopg2
import pandas as pd
import streamlit as st
import matplotlib.pyplot as plt
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

# --- PALETTE (cohérente avec la version desktop "Sobre & Institutionnel") ---
COULEUR_SIDEBAR = "#1B2A41"
COULEUR_ACCENT = "#2E5EAA"
COULEUR_AVERTISSEMENT = "#D68910"

st.set_page_config(page_title="Gestion de Tontine", page_icon="🧵", layout="wide")

st.markdown(f"""
<style>
[data-testid="stSidebar"] {{ background-color: {COULEUR_SIDEBAR}; }}
[data-testid="stSidebar"] * {{ color: #FFFFFF !important; }}
[data-testid="stSidebar"] button {{ background-color: transparent; border: 1px solid #3B5171; text-align:left; }}
[data-testid="stSidebar"] button:hover {{ background-color: #24344A; border-color: #24344A; }}
div[data-testid="stMetric"] {{ background-color: #FFFFFF; border: 1px solid #D7DEE8; border-radius: 10px; padding: 12px; }}
</style>
""", unsafe_allow_html=True)

# ==========================================
# CONNEXION BASE DE DONNÉES (PostgreSQL)
# ==========================================
def get_connection():
    return psycopg2.connect(st.secrets["DATABASE_URL"])

@st.cache_resource
def get_cached_connection():
    return get_connection()

def run_query(sql, params=None, fetch=False, fetchone=False):
    conn = get_cached_connection()
    try:
        with conn:
            with conn.cursor() as cur:
                cur.execute(sql, params or ())
                if fetchone:
                    return cur.fetchone()
                if fetch:
                    return cur.fetchall()
                return None
    except (psycopg2.InterfaceError, psycopg2.OperationalError):
        # La connexion a été coupée (inactivité, redémarrage réseau...) : on la recrée une fois
        get_cached_connection.clear()
        conn = get_cached_connection()
        with conn:
            with conn.cursor() as cur:
                cur.execute(sql, params or ())
                if fetchone:
                    return cur.fetchone()
                if fetch:
                    return cur.fetchall()
                return None

def init_db():
    run_query('''CREATE TABLE IF NOT EXISTS types_pagne (
        id SERIAL PRIMARY KEY, nom TEXT NOT NULL, prix_unitaire REAL NOT NULL)''')
    run_query('''CREATE TABLE IF NOT EXISTS membres (
        id SERIAL PRIMARY KEY, nom TEXT NOT NULL, telephone TEXT,
        objectif_financier REAL DEFAULT 0.0, type_pagne_id INTEGER REFERENCES types_pagne(id))''')
    run_query('''CREATE TABLE IF NOT EXISTS versements (
        id SERIAL PRIMARY KEY, membre_id INTEGER REFERENCES membres(id),
        montant REAL NOT NULL, date_versement TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
        enregistre_par TEXT)''')
    run_query('''CREATE TABLE IF NOT EXISTS journal_actions (
        id SERIAL PRIMARY KEY, date_action TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
        operateur TEXT NOT NULL, action TEXT NOT NULL, details TEXT)''')

init_db()

# ==========================================
# OUTILS COMMUNS
# ==========================================
def recuperer_types_pagne():
    return run_query("SELECT id, nom, prix_unitaire FROM types_pagne ORDER BY nom ASC", fetch=True)

def options_types_dict(types):
    options = {"Non défini": None}
    for tid, nom, prix in types:
        options[f"{nom} — {prix:,.0f} FCFA".replace(",", " ")] = tid
    return options

def enregistrer_action(operateur, action, details=""):
    run_query("INSERT INTO journal_actions (operateur, action, details) VALUES (%s, %s, %s)", (operateur, action, details))

# ==========================================
# PAGE : TABLEAU DE BORD
# ==========================================
def afficher_accueil():
    st.title("📊 Tableau de Bord & Évolution")

    total_membres = run_query("SELECT COUNT(id) FROM membres", fetch=True)[0][0]
    total_caisse = run_query("SELECT COALESCE(SUM(montant), 0) FROM versements", fetch=True)[0][0]

    col1, col2 = st.columns(2)
    col1.metric("👥 Participants", total_membres)
    col2.metric("💰 Total en caisse", f"{total_caisse:,.0f} FCFA".replace(",", " "))

    st.markdown("#### Évolution cumulée de l'épargne")
    evolution = run_query(
        "SELECT DATE(date_versement) d, SUM(montant) m FROM versements GROUP BY DATE(date_versement) ORDER BY d ASC",
        fetch=True
    )
    if evolution:
        dates = [str(r[0]) for r in evolution]
        cumul, total_temp = [], 0
        for _, m in evolution:
            total_temp += float(m)
            cumul.append(total_temp)
        fig, ax = plt.subplots(figsize=(9, 3))
        ax.bar(dates, cumul, color=COULEUR_ACCENT)
        ax.set_ylabel("FCFA")
        ax.grid(True, axis='y', linestyle='--', alpha=0.4)
        plt.xticks(rotation=30, ha="right")
        st.pyplot(fig)
    else:
        st.info("Aucune donnée de versement disponible.")

    st.markdown("#### Cotisations par qualité de pagne")
    repartition = run_query('''
        SELECT COALESCE(tp.nom, 'Non défini') AS nom, COALESCE(SUM(v.montant), 0) AS total
        FROM membres m
        LEFT JOIN versements v ON v.membre_id = m.id
        LEFT JOIN types_pagne tp ON m.type_pagne_id = tp.id
        GROUP BY tp.id, tp.nom
    ''', fetch=True)
    valides = [(n, t) for n, t in repartition if t and t > 0]
    if valides:
        noms = [v[0] for v in valides]
        totaux = [v[1] for v in valides]
        palette = ["#2E5EAA", "#4E7FC7", "#1B3A63", "#6FA0DE", "#0F2A4A", "#8CB6E8"]
        couleurs, i = [], 0
        for n in noms:
            if n == "Non défini":
                couleurs.append(COULEUR_AVERTISSEMENT)
            else:
                couleurs.append(palette[i % len(palette)])
                i += 1
        fig2, ax2 = plt.subplots(figsize=(5, 5))
        ax2.pie(totaux, labels=noms, colors=couleurs, autopct=lambda p: f"{p:.0f}%" if p >= 5 else "", startangle=90)
        ax2.axis("equal")
        st.pyplot(fig2)
    else:
        st.info("Aucune cotisation enregistrée pour le moment.")

# ==========================================
# PAGE : GESTION DES MEMBRES
# ==========================================
def afficher_membres():
    st.title("👥 Gestion des Membres")

    types = recuperer_types_pagne()
    options = options_types_dict(types)

    with st.form("form_ajout_membre", clear_on_submit=True):
        c1, c2, c3, c4 = st.columns(4)
        nom = c1.text_input("Nom et Prénom")
        tel = c2.text_input("Téléphone")
        objectif = c3.text_input("Objectif (FCFA)")
        type_choisi = c4.selectbox("Qualité de pagne", list(options.keys()))
        operateur = st.text_input("Qui enregistre cet ajout ?")
        if st.form_submit_button("Ajouter le membre"):
            if not nom or not objectif:
                st.error("Le nom et l'objectif sont obligatoires.")
            elif not operateur.strip():
                st.error("Le nom de l'opérateur est requis.")
            else:
                try:
                    obj_f = float(objectif)
                except ValueError:
                    st.error("L'objectif doit être un nombre.")
                    st.stop()
                run_query("INSERT INTO membres (nom, telephone, objectif_financier, type_pagne_id) VALUES (%s, %s, %s, %s)",
                           (nom, tel, obj_f, options[type_choisi]))
                enregistrer_action(operateur.strip(), "Ajout membre", f"Membre « {nom} » ajouté (objectif {obj_f:,.0f} FCFA)".replace(",", " "))
                st.success("Membre ajouté.")
                st.rerun()

    if not types:
        st.warning("⚠️ Aucun type de pagne défini — configure-les dans « Types de Pagnes ».")

    st.markdown("---")
    membres = run_query('''SELECT m.id, m.nom, m.telephone, m.objectif_financier, m.type_pagne_id, tp.nom
                            FROM membres m LEFT JOIN types_pagne tp ON m.type_pagne_id = tp.id
                            ORDER BY m.id DESC''', fetch=True)

    for mid, nom, tel, obj, type_id_actuel, type_nom in membres:
        with st.expander(f"{nom} — {type_nom or 'Type non défini'}"):
            st.write(f"Téléphone : {tel or '-'}  |  Objectif : {obj:,.0f} FCFA".replace(",", " "))
            colA, colB = st.columns(2)

            with colA:
                with st.form(f"edit_{mid}"):
                    st.markdown("**Modifier**")
                    nouveau_nom = st.text_input("Nom", value=nom, key=f"nom_{mid}")
                    nouveau_tel = st.text_input("Téléphone", value=tel or "", key=f"tel_{mid}")
                    nouvel_obj = st.text_input("Objectif", value=str(obj), key=f"obj_{mid}")

                    types_actuels = recuperer_types_pagne()
                    opts = options_types_dict(types_actuels)
                    labels = list(opts.keys())
                    label_actuel = next((lbl for lbl, tid in opts.items() if tid == type_id_actuel), "Non défini")
                    index_actuel = labels.index(label_actuel) if label_actuel in labels else 0
                    nouveau_type = st.selectbox("Qualité de pagne", labels, index=index_actuel, key=f"type_{mid}")

                    op_edit = st.text_input("Opérateur", key=f"opedit_{mid}")
                    if st.form_submit_button("Enregistrer les modifications"):
                        if not op_edit.strip():
                            st.error("Nom de l'opérateur requis.")
                        else:
                            try:
                                nv_obj_f = float(nouvel_obj)
                            except ValueError:
                                st.error("Objectif invalide.")
                                st.stop()
                            run_query("UPDATE membres SET nom=%s, telephone=%s, objectif_financier=%s, type_pagne_id=%s WHERE id=%s",
                                       (nouveau_nom, nouveau_tel, nv_obj_f, opts[nouveau_type], mid))
                            enregistrer_action(op_edit.strip(), "Modification membre", f"« {nom} » modifié -> nom: {nouveau_nom}, objectif: {nv_obj_f:,.0f} FCFA".replace(",", " "))
                            st.success("Membre modifié.")
                            st.rerun()

            with colB:
                with st.form(f"delete_{mid}"):
                    st.markdown("**Supprimer**")
                    op_del = st.text_input("Opérateur (confirme la suppression)", key=f"opdel_{mid}")
                    if st.form_submit_button("❌ Supprimer ce membre"):
                        if not op_del.strip():
                            st.error("Nom de l'opérateur requis.")
                        else:
                            run_query("DELETE FROM versements WHERE membre_id=%s", (mid,))
                            run_query("DELETE FROM membres WHERE id=%s", (mid,))
                            enregistrer_action(op_del.strip(), "Suppression membre", f"Membre « {nom} » (ID {mid}) supprimé, avec son historique de versements")
                            st.success("Membre supprimé.")
                            st.rerun()

# ==========================================
# PAGE : SAISIE DES VERSEMENTS
# ==========================================
def afficher_versements():
    st.title("💰 Saisir un Versement")

    membres = run_query("SELECT id, nom FROM membres ORDER BY nom ASC", fetch=True)
    if not membres:
        st.warning("Ajoute d'abord des membres.")
        return

    options_membres = {f"{nom} (ID {mid})": mid for mid, nom in membres}
    with st.form("form_versement", clear_on_submit=True):
        c1, c2, c3 = st.columns(3)
        membre_choisi = c1.selectbox("Participant", list(options_membres.keys()))
        montant = c2.text_input("Montant (FCFA)")
        operateur = c3.text_input("Opérateur")
        if st.form_submit_button("Valider le versement"):
            if not operateur.strip():
                st.error("Nom de l'opérateur requis.")
            else:
                try:
                    m = float(montant)
                    if m <= 0:
                        raise ValueError
                except ValueError:
                    st.error("Le montant doit être un nombre positif.")
                    st.stop()
                mid = options_membres[membre_choisi]
                run_query("INSERT INTO versements (membre_id, montant, enregistre_par) VALUES (%s, %s, %s)", (mid, m, operateur.strip()))
                nom_membre = membre_choisi.split(" (ID")[0]
                enregistrer_action(operateur.strip(), "Versement", f"{m:,.0f} FCFA reçus de « {nom_membre} »".replace(",", " "))
                st.success("Versement enregistré.")
                st.rerun()

    st.markdown("---")
    st.subheader("Suivi des objectifs")
    suivi = run_query('''SELECT m.nom, m.objectif_financier, COALESCE(SUM(v.montant), 0)
                          FROM membres m LEFT JOIN versements v ON v.membre_id = m.id
                          GROUP BY m.id, m.nom, m.objectif_financier ORDER BY m.nom''', fetch=True)
    for nom, obj, total in suivi:
        prog = min(total / obj, 1.0) if obj > 0 else 0
        st.write(f"**{nom}** — {total:,.0f} / {obj:,.0f} FCFA".replace(",", " "))
        st.progress(prog)

# ==========================================
# PAGE : FICHE DE POINTAGE
# ==========================================
def afficher_pointage():
    st.title("📋 Fiche de Pointage")
    st.caption("Registre complet des versements avec horodatage et opérateur.")

    rows = run_query('''SELECT v.id, v.date_versement, m.nom, v.montant, v.enregistre_par
                         FROM versements v JOIN membres m ON v.membre_id = m.id
                         ORDER BY v.date_versement DESC''', fetch=True)
    if not rows:
        st.info("Aucun versement enregistré pour le moment.")
        return

    df = pd.DataFrame(rows, columns=["Réf", "Date", "Participant", "Montant (FCFA)", "Enregistré par"])
    st.dataframe(df, use_container_width=True, hide_index=True)

# ==========================================
# PAGE : TYPES DE PAGNES
# ==========================================
def afficher_types_pagne():
    st.title("🧵 Gestion des Types de Pagnes")

    with st.form("form_type", clear_on_submit=True):
        c1, c2 = st.columns(2)
        nom = c1.text_input("Nom de la qualité (ex: Wax Hollandais)")
        prix = c2.text_input("Prix unitaire (FCFA)")
        if st.form_submit_button("Ajouter"):
            if not nom or not prix:
                st.error("Le nom et le prix sont obligatoires.")
            else:
                try:
                    p = float(prix)
                except ValueError:
                    st.error("Le prix doit être un nombre.")
                    st.stop()
                run_query("INSERT INTO types_pagne (nom, prix_unitaire) VALUES (%s, %s)", (nom, p))
                st.success("Type ajouté.")
                st.rerun()

    st.markdown("---")
    types = run_query("SELECT id, nom, prix_unitaire FROM types_pagne ORDER BY id DESC", fetch=True)
    for tid, nom, prix in types:
        nb = run_query("SELECT COUNT(id) FROM membres WHERE type_pagne_id=%s", (tid,), fetch=True)[0][0]
        c1, c2, c3, c4 = st.columns([3, 2, 2, 2])
        c1.write(nom)
        c2.write(f"{prix:,.0f} FCFA".replace(",", " "))
        c3.write(f"{nb} membre(s)")
        if c4.button("❌ Supprimer", key=f"deltype_{tid}"):
            if nb > 0:
                st.error(f"Ce type est utilisé par {nb} membre(s). Réaffecte-les avant de le supprimer.")
            else:
                run_query("DELETE FROM types_pagne WHERE id=%s", (tid,))
                st.rerun()

# ==========================================
# PAGE : JOURNAL DES ACTIONS
# ==========================================
def afficher_journal():
    st.title("🕵️ Journal des Actions")
    st.caption("Historique de qui a fait quoi : versements et opérations sur les membres.")

    filtre = st.selectbox("Filtrer par type d'action", ["Toutes", "Versement", "Ajout membre", "Modification membre", "Suppression membre"])
    if filtre == "Toutes":
        rows = run_query("SELECT date_action, operateur, action, details FROM journal_actions ORDER BY date_action DESC", fetch=True)
    else:
        rows = run_query("SELECT date_action, operateur, action, details FROM journal_actions WHERE action=%s ORDER BY date_action DESC", (filtre,), fetch=True)

    if not rows:
        st.info("Aucune action enregistrée pour ce filtre.")
        return

    df = pd.DataFrame(rows, columns=["Date", "Opérateur", "Action", "Détails"])
    st.dataframe(df, use_container_width=True, hide_index=True)

# ==========================================
# PAGE : CLÔTURE & RELIQUATS
# ==========================================
def afficher_cloture():
    st.title("📊 Clôture & Reliquats")
    st.caption("Le reliquat de chaque participant est calculé selon la qualité de pagne qui lui est assignée.")

    if st.button("Calculer le bilan"):
        st.session_state["bilan"] = run_query('''
            SELECT m.nom, tp.nom, tp.prix_unitaire, COALESCE(SUM(v.montant), 0)
            FROM membres m
            LEFT JOIN versements v ON v.membre_id = m.id
            LEFT JOIN types_pagne tp ON m.type_pagne_id = tp.id
            GROUP BY m.id, m.nom, tp.nom, tp.prix_unitaire
            ORDER BY m.nom ASC
        ''', fetch=True)

    bilan = st.session_state.get("bilan")
    if bilan:
        data = []
        for nom, type_nom, prix, total in bilan:
            if type_nom is None:
                data.append([nom, "Non défini", "-", f"{total:,.0f}".replace(",", " "), "-", "Type non défini"])
            else:
                reliquat = total - prix
                statut = "Trop-perçu" if reliquat > 0 else ("Équilibré" if reliquat == 0 else "Déficitaire")
                data.append([nom, type_nom, f"{prix:,.0f}".replace(",", " "), f"{total:,.0f}".replace(",", " "),
                             f"{reliquat:,.0f}".replace(",", " "), statut])
        df = pd.DataFrame(data, columns=["Participant", "Qualité", "Prix pagne", "Total cotisé", "Reliquat", "Statut"])
        st.dataframe(df, use_container_width=True, hide_index=True)

        if st.button("📄 Générer le rapport Word"):
            doc = Document()
            doc.add_heading('Bilan - Opération Pagnes', 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph("Le tableau ci-dessous détaille la situation comptable de chaque participant selon la qualité de pagne qui lui a été assignée.")
            table = doc.add_table(rows=1, cols=5)
            table.style = 'Table Grid'
            hdr = table.rows[0].cells
            for i, h in enumerate(['Participant', 'Qualité de Pagne', 'Prix Pagne (FCFA)', 'Total Cotisé (FCFA)', 'Reliquat / Reste (FCFA)']):
                hdr[i].text = h
            total_caisse = 0
            for nom, type_nom, prix, total in bilan:
                total_caisse += total
                row = table.add_row().cells
                row[0].text = str(nom)
                row[1].text = type_nom if type_nom else "Non défini"
                row[2].text = f"{prix:,.0f}" if prix is not None else "-"
                row[3].text = f"{total:,.0f}"
                row[4].text = f"{total - prix:,.0f}" if prix is not None else "-"
            doc.add_paragraph(f"\nLa trésorerie globale de ce cycle s'élève à {total_caisse:,.0f} FCFA.")
            buf = io.BytesIO()
            doc.save(buf)
            buf.seek(0)
            st.download_button("Télécharger le rapport Word", buf, file_name="Bilan_Tontine.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    else:
        st.info("Clique sur « Calculer le bilan » pour afficher les résultats.")

# ==========================================
# NAVIGATION
# ==========================================
PAGES = {
    "🏠 Tableau de bord": afficher_accueil,
    "👥 Gestion des Membres": afficher_membres,
    "💰 Saisir un Versement": afficher_versements,
    "📋 Fiche de Pointage": afficher_pointage,
    "🧵 Types de Pagnes": afficher_types_pagne,
    "🕵️ Journal des Actions": afficher_journal,
    "📊 Clôture & Reliquats": afficher_cloture,
}

if "page" not in st.session_state:
    st.session_state.page = list(PAGES.keys())[0]

with st.sidebar:
    st.markdown("## Tontine App")
    for nom_page in PAGES:
        if st.button(nom_page, use_container_width=True, key=f"nav_{nom_page}"):
            st.session_state.page = nom_page

PAGES[st.session_state.page]()
